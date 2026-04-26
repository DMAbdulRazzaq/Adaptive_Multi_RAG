"""Stage 1 — Document Ingestion"""
import os, re, hashlib
from dataclasses import dataclass, field
from datetime import datetime

@dataclass
class RawDocument:
    document_id: str; file_path: str; file_type: str; title: str; text: str
    metadata: dict = field(default_factory=dict)
    page_count: int = 0; word_count: int = 0
    loaded_at: str = field(default_factory=lambda: datetime.utcnow().isoformat())

class DocumentLoader:
    def load(self, file_path):
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Not found: {file_path}")
        ext = os.path.splitext(file_path)[1].lower().lstrip(".")
        loaders = {"pdf":self._load_pdf,"txt":self._load_txt,"docx":self._load_docx,"md":self._load_txt}
        if ext not in loaders:
            raise ValueError(f"Unsupported: .{ext}")
        raw, meta = loaders[ext](file_path)
        clean = self._clean(raw)
        return RawDocument(
            document_id=self._id(file_path), file_path=file_path, file_type=ext,
            title=meta.get("title") or os.path.splitext(os.path.basename(file_path))[0],
            text=clean, metadata=meta, page_count=meta.get("page_count",0),
            word_count=len(clean.split()),
        )
    def load_directory(self, dir_path):
        supported = {".pdf",".txt",".docx",".md"}
        docs = []
        for f in os.listdir(dir_path):
            if os.path.splitext(f)[1].lower() in supported:
                try: docs.append(self.load(os.path.join(dir_path, f)))
                except Exception as e: print(f"[Skip] {f}: {e}")
        return docs
    def _load_pdf(self, path):
        try:
            import fitz
            doc = fitz.open(path)
            meta = doc.metadata or {}
            text = "\n\n".join(p.get_text("text") for p in doc)
            doc.close()
            return text, {"title":meta.get("title",""),"page_count":doc.page_count,"source":"pymupdf"}
        except ImportError:
            try:
                import pdfplumber
                with pdfplumber.open(path) as pdf:
                    return "\n\n".join(p.extract_text() or "" for p in pdf.pages), {"page_count":len(pdf.pages),"source":"pdfplumber"}
            except ImportError:
                raise ImportError("pip install pymupdf")
    def _load_txt(self, path):
        text = open(path,"r",encoding="utf-8",errors="replace").read()
        title = next((l.strip("# ").strip() for l in text.splitlines() if l.strip()), "")
        return text, {"title":title,"page_count":1,"source":"plaintext"}
    def _load_docx(self, path):
        try:
            from docx import Document
        except ImportError:
            raise ImportError("pip install python-docx")
        doc = Document(path)
        text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        c = doc.core_properties
        return text, {"title":c.title or "","page_count":len(doc.sections),"source":"python-docx"}
    def _clean(self, text):
        text = re.sub(r"[^\x09\x0A\x20-\x7E\u00A0-\uFFFF]"," ",text)
        text = re.sub(r"[ \t]+"," ",text)
        text = re.sub(r"^\s*\d+\s*$","",text,flags=re.MULTILINE)
        text = re.sub(r"\n{3,}","\n\n",text)
        return text.strip()
    @staticmethod
    def _id(path): return hashlib.md5(os.path.abspath(path).encode()).hexdigest()[:12]